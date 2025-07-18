import streamlit as st
from llama_index.core.prompts import RichPromptTemplate
from llama_index.llms.ollama import Ollama
from llama_index.llms.openai import OpenAI
from llama_index.core.llms import ChatMessage, TextBlock, ImageBlock
from llama_index.embeddings.ollama import OllamaEmbedding
from llama_index.embeddings.openai import OpenAIEmbedding
import time
import os
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import WebDriverException
import logging
from llama_index.core.callbacks import LlamaDebugHandler, CallbackManager
from llama_index.core.agent.workflow import FunctionAgent
from llama_index.core.tools import FunctionTool
from llama_index.core.settings import Settings
from pathlib import Path
from git import Repo
import ollama
from llama_index.core import (
    VectorStoreIndex,
    SimpleDirectoryReader,
    StorageContext,
    load_index_from_storage,
    set_global_handler
)

import phoenix as px
px.launch_app()
set_global_handler("arize_phoenix")

# import logging
# import sys

# logging.basicConfig(stream=sys.stdout, level=logging.DEBUG)
# logging.getLogger().addHandler(logging.StreamHandler(stream=sys.stdout))


# import dotenv

# # Load environment variables
# dotenv.load_dotenv()
# import debugpy

# if __name__ == "__main__":
#     debugpy.listen(("localhost", 5678))
#     print("Waiting for debugger attach...")
#     debugpy.wait_for_client()
#     print("Debugger attached")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

debug_handler = LlamaDebugHandler(print_trace_on_end=True)
callback_manager = CallbackManager([debug_handler])

# Set page config
st.set_page_config(
    page_title="AI-Powered Testing Assistant",
    page_icon="🧪",
    layout="wide"
)

# Debugging toggle
DEBUG = st.sidebar.checkbox("Enable Debug Mode", value=False)

OPENAI_API_KEY=""
embedProvider=""

def log_debug(message):
    if DEBUG:
        logger.info(f"DEBUG: {message}")
        st.sidebar.text(f"DEBUG: {message}")

# Sidebar configuration
with st.sidebar:
    st.title("Configuration")
    model_provider = st.selectbox(
        "Select Model Provider",
        ["Ollama (Local)", "OpenAI"],
        index=0
    )
    local_models=[i["model"] for i in ollama.list()["models"]]
    try:
        modelid=local_models.index("gemma3:12b")
    except:
        modelid=[ollama.list().models.index(id) for id in  ollama.list().models if 'M' not in id.details.parameter_size][0]
        #modelid=0
        log_debug("gemma3:12b model not found, using first model in list")
    if model_provider == "Ollama (Local)":
        model_name = st.selectbox(
        "Select Model Name",
        local_models,
        index=modelid
        )
        index=0 
        log_debug(f"Selected model name {model_name}")
        #model_name=st.text_input("Model name", "gemma3:12b")
        thinking=False
        if "qwen" in model_name or "deepseek" in model_name:
            thinking=st.checkbox("Enable Thinking",value=False)
            log_debug(f"Setting Thinking {thinking}")
        temperature = st.slider("Temperature (Creativity)", 0.0, 1.0, 0.1)
        context_window = st.number_input("Context Window", 4096, 32768, 16384)
    else:
        if not os.getenv("OPENAI_API_KEY"):
            OPENAI_API_KEY=st.text_input(
                "OpenAI API Key",
                placeholder="Please Enter your OpenAI API key")
        else:
            OPENAI_API_KEY=os.getenv("OPENAI_API_KEY")
        model_name = st.text_input("Model name", "gpt-4.1-mini")
        temperature = st.slider("Temperature", 0.0, 1.0, 0.1)
        max_tokens = st.number_input("Max Tokens", 512, 8192, 4096)
    
    st.markdown("---")

# Initialize LLM based on selection
def get_llm():
    if model_provider == "Ollama (Local)":
        #log_debug("Model Thinking set to: "+ str(thinking))
        if thinking:
            return Ollama(
                model=model_name,
                temperature=temperature,
                context_window=context_window,
                request_timeout=500,
                thinking=thinking
            )
        else:
            return Ollama(
                model=model_name,
                temperature=temperature,
                context_window=context_window,
                request_timeout=500
            )
    else:
        if OPENAI_API_KEY:
            return OpenAI(
                model=model_name,
                temperature=temperature,
                max_tokens=max_tokens,
                request_timeout=500,
                api_key=OPENAI_API_KEY
            )
        else:
            return OpenAI(
                model=model_name,
                temperature=temperature,
                max_tokens=max_tokens,
                request_timeout=500
            )

def get_ollamaEmbeddings():
    embed_model=OllamaEmbedding(
    model_name="mxbai-embed-large:latest",
    embed_batch_size=32
    )
    return embed_model

def get_openaiEmbeddings():
    if OPENAI_API_KEY:
        embed_model=OpenAIEmbedding(
            model_name="text-embedding-3-small",
            api_key=OPENAI_API_KEY
        )
    else:
        embed_model=OpenAIEmbedding(
            model_name="text-embedding-3-small"
        )
    return embed_model

def set_ollama():
    Settings.llm=get_llm()
    Settings.embed_model=get_ollamaEmbeddings()
    

# Initialize or load index

def _load_init_index(repo_path,embProvider):
    """Load existing index"""
    index_path=f"{repo_path}/{embProvider}/.code_index"
    try:
        storage_context = StorageContext.from_defaults(persist_dir=str(index_path))
        index = load_index_from_storage(storage_context)
    except:
    # Load all Python files
        documents = SimpleDirectoryReader(
            input_dir=str(repo_path),
            required_exts=[".py"],
            recursive=True
        ).load_data()
        
        # Create and save index
        index = VectorStoreIndex.from_documents(documents)
        index.storage_context.persist(persist_dir=str(index_path))
    return index

def _create_agent():
    """Create a ReAct agent with custom tools"""
    # Define tools
    check_changes_tool = FunctionTool.from_defaults(fn=check_code_changes)
    generate_test_cases_tool = FunctionTool.from_defaults(fn=generate_test_cases)
    update_index_tool = FunctionTool.from_defaults(fn=_load_init_index)
    generate_tests_tool = FunctionTool.from_defaults(fn=generate_tests)
    
    agent = FunctionAgent(
        tools=[check_changes_tool,generate_test_cases_tool,generate_tests_tool,update_index_tool],
        llm=Settings.llm,
        verbose=True,
        CallbackManager=CallbackManager
    )

    return agent

def check_code_changes(repo_path:str,since_commit: str = "HEAD~1",agentic=True) -> str:
    """
    Check what Python files have changed since a specific commit.
    
    Args:
        since_commit: The commit to compare against (default: previous commit)
        
    Returns:
        String describing the changed files and their diffs
    """
    
    try:
        repo = Repo(repo_path)
        print(f"Repository: {repo_path}")
        print(f"Comparing: {since_commit} ({repo.commit(since_commit).hexsha[:7]}) to HEAD ({repo.head.commit.hexsha[:7]})")
    except Exception as e:
        print(f"Error opening repository at {repo_path}: {str(e)}")

    # Get changed files
    changed_files = []
    diffs = repo.commit(since_commit).diff(repo.head.commit)
    if not diffs:
        print("No differences found between the commits")
    else:
        for diff in diffs:
            file_path = getattr(diff, "a_path", None) or getattr(diff, "b_path", None)
            # Some diffs may not have a diff attribute (e.g., binary files or renames)
            if ".code_index" in file_path:
                continue
            diff_content = ""
            if hasattr(diff, "diff") and diff.diff:
                try:
                    diff_content = diff.diff.decode("utf-8", errors="replace")
                except Exception:
                    diff_content = "<Could not decode diff>"
            print(f"Found changes in {file_path} ({getattr(diff, 'change_type', '?')})")
            if diff_content:
                print(f"Diff length: {len(diff_content)} bytes")
            # If diff_content is empty, try to get the diff using repo.git.diff as a fallback
            if not diff_content and file_path:
                try:
                    diff_content = repo.git.diff(f"{since_commit}..HEAD", "--", file_path)
                except Exception:
                    diff_content = "<Could not retrieve diff via git command>"
            changed_files.append({
                "path": file_path,
                "change_type": getattr(diff, "change_type", "?"),
                "diff": diff_content,
            })

    # Format output
    if agentic:
        output = ["Changed Python files:"]
        for file in changed_files:
            output.append(f"\nFile: {file['path']} ({file['change_type']})")
            output.append(f"Diff:\n{file['diff']}")
        return "\n".join(output)
    else:
        output = []
        for file in changed_files:
            output.append(f"{file['path']}")
        return output

def generate_tests( repo_path,file_path: str, framework: str = "pytest",streaming=False) -> str:
    """
    Generate tests for a specific Python file.
    
    Args:
        file_path: Path to the Python file
        framework: Testing framework to use (default: pytest)
        
    Returns:
        Generated test code as a string
    """
    # Read the file content
    if repo_path=="":
        return "Error: Repository path is not specified."
    repopath=Path(repo_path)
    full_path = repopath / file_path
    if not full_path.exists():
        return f"Error: File {file_path} does not exist."
        
    with open(full_path, "r") as f:
        code_content = f.read()
    
    # Use LLM to generate tests
    prompt = f"""
    Analyze the following Python code and generate comprehensive test cases using {framework}.
    Focus on edge cases, input validation, and expected behavior.
    
    Code:
    {code_content}
    
    Please return only the test code with appropriate imports, no additional explanation.
    The tests should be production-ready and follow best practices.
    """
    llm=Settings.llm
    if streaming:
        response=llm.stream_complete(prompt)
        return response
    else:
        response=llm.complete(prompt)
        return response



def generate_test_cases(index,changes: str,streaming=False) -> str:
    """
    Generate test cases for a specific Python file.
    
    Args:
        file_path: Path to the Python file
        
    Returns:
        Generated test cases as a markup format
    """
    
    # Use LLM to generate tests
    prompt = f"""
        You are an expert test engineer who excels in creating test cases for IT products, 
        Analyze the following Python code and create test cases, Focusing on edge cases, input validation, and expected behavior.
        
        Code:
        {changes}
        
        The test cases should follow the following markdown format: 
"| Test Case ID | TC-PERF-001 |\n",
"|--------------|------------|\n",
"| **Title** | Verify API response time under load |\n",
"| **Description** | Ensure GraphQL API responds within 500ms under expected load |\n",
"| **Preconditions** | Load testing environment set up |\n",
"| **Test Steps** | 1. Simulate concurrent users accessing API endpoints<br>2. Measure response times |\n",
"| **Expected Result** | API response times remain below 500ms under load |\n",
"\n",
"---\n",
"\n",
""" 
    qe=index.as_query_engine(streaming=streaming, similarity_top_k=1)
    if streaming:
        return qe.query(prompt).response_gen
    else:
        return qe.query(prompt).response
        
# Configure Chrome options
def get_chrome_options():
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    return chrome_options

# Main app
#"👨‍✈️🕵🧙🧙‍♂️👨‍💻👩‍💻👨‍⚕️🧑‍🔬👨‍🎓🧞🧞‍♀️🧞‍♂️🧜‍♀️🧜‍♂️🧚‍♀️🧚‍♂️🧚🧙‍♂️"
st.title("🧙‍♂️ AI-Powered Testing Magician⚡")
st.markdown("""
This application demonstrates three AI-powered testing capabilities:
1. **Test Case Generation** from BRD/solution documents
2. **Visual Test Case Generation** from screenshots
3. **AI-Driven Test Data Generation** for testing scenarios
4. **AI-Driven Test Cases Generation** for code changes.
5. **AI-Driven Unit Tests Generation** for code changes.
""")

# Tab layout
tab1, tab2, tab3, tab4, tab5= st.tabs([
    "Test Case Generation", 
    "Visual Test Case Generation", 
    "Test Data Generation",
    "Code Change Test Cases Generation",
    "Code Change Unit Tests Generation"
])

if 'llm_response' not in st.session_state:
    st.session_state.llm_response = None
    
# Tab 1: Test Case Generation from BRD
with tab1:
    st.header("Generate Test Cases from BRD Document")
    
    # Create two columns for upload and text area
    col1, col2 = st.columns(2)
    
    with col1:
        uploaded_file = st.file_uploader(
            "Upload BRD Document",
            type=["txt", "pdf", "docx", "doc"],
            help="Upload your Business Requirements Document (supports .txt, .pdf, .docx)"
        )
        try:
        # If file is uploaded, extract text
            if uploaded_file:
                file_type = uploaded_file.type
                if "pdf" in file_type:
                    st.markdown("Extracting text from PDF is not supported, please use text file or docx...")
                    brd_text=""
                elif "word" in file_type or "doc" in file_type.lower():
                    # # Handle Word document extraction (would need python-docx)
                    # import docx
                    # doc = docx.Document(uploaded_file)
                    #brd_text = "\n".join([para.text for para in doc.paragraphs])
                    from llama_index.readers.file import DocxReader
                    with open("uploadfile.docx","wb") as f:
                        f.writelines(uploaded_file.readlines())
                    doc = DocxReader().load_data(file="uploadfile.docx")
                    brd_text = doc[0].text
                    os.remove("uploadfile.docx")
                else:  # Assume plain text
                    brd_text = uploaded_file.getvalue().decode("utf-8")
                st.markdown(brd_text)
        except Exception as e:
            st.error(f"Error generating test cases: {str(e)}")
            log_debug(f"Error details: {e}")
    
    with col2:
        brd_text = st.text_area(
            "Or paste your BRD document content here:",
            height=300,
            help="Business Requirements Document content to generate test cases from"
        )
    
    if st.button("Generate Test Cases"):
        # Check if either file is uploaded or text is pasted
        if not brd_text and not uploaded_file:
            st.warning("Please upload a BRD document or paste the content first")
        else:
            with st.spinner("Generating test cases..."):
                try:
                    # If file is uploaded, extract text
                    if uploaded_file:
                        file_type = uploaded_file.type
                        if "pdf" in file_type:
                            # Handle PDF extraction (would need PyPDF2 or similar)
                            import PyPDF2
                            pdf_reader = PyPDF2.PdfReader(uploaded_file)
                            brd_text = ""
                            for page in pdf_reader.pages:
                                brd_text += page.extract_text()
                        elif "word" in file_type or "doc" in file_type.lower():
                            # Handle Word document extraction (would need python-docx)
                            import docx
                            doc = docx.Document(uploaded_file)
                            brd_text = "\n".join([para.text for para in doc.paragraphs])
                        else:  # Assume plain text
                            brd_text = uploaded_file.getvalue().decode("utf-8")
                    
                    # Build the prompt
                    template_str = """You are an expert test engineer who excels in creating test cases for web products, create test cases from the following Business Requirements Document (BRD):
                    ---------------------
                    {{ context_str }}
                    ---------------------
                    """
                    qa_template = RichPromptTemplate(template_str)
                    messages = qa_template.format_messages(context_str=brd_text)
                    
                    log_debug(f"Prompt messages: {messages}")
                    
                    # Get LLM response
                    llm = get_llm()
                    gen = llm.stream_chat(messages=messages)
                    
                    # Stream the output
                    output_container = st.empty()
                    full_response = ""
                    
                    for response in gen:
                        full_response += response.delta
                        output_container.markdown(full_response)
                        time.sleep(0.01)
                    st.session_state['llm_response'] = full_response
                    
                    datetime_stamp = f"{datetime.now().strftime('%Y-%m-%d-%H-%M-%S')}"
                    st.download_button(
                        label="Export Generated Test Cases",
                        data=full_response,
                        file_name=f'test_cases{datetime_stamp}.md',
                        mime='text/md')
                    
                except Exception as e:
                    st.error(f"Error generating test cases: {str(e)}")
                    log_debug(f"Error details: {e}")

# Tab 2: Visual Test Case Generation
with tab2:
    st.header("Generate Test Cases from Screenshots")
    
    option = st.radio(
        "Select screenshot source:",
        ("Upload an image", "Capture from URL")
    )
    
    image_path = None
    
    if option == "Upload an image":
        uploaded_file = st.file_uploader(
            "Upload a screenshot",
            type=["png", "jpg", "jpeg"]
        )
        if uploaded_file:
            # Save the uploaded file
            image_path = "uploaded_screenshot.png"
            with open(image_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Display the image
            st.image(uploaded_file, caption="Uploaded Screenshot")
    
    if option=="Capture from URL":  # Capture from URL
        url = st.text_input("Enter URL to capture:", placeholder="https://example.com")
        if st.button("Capture Screenshot"):
            if not url:
                st.warning("Please enter a URL first")
            else:
                with st.spinner("Capturing screenshot..."):
                    try:
                        image_path = "captured_screenshot.png"
                        
                        log_debug(f"Attempting to capture URL: {url}")
                        
                        # Configure Chrome with additional options
                        chrome_options = get_chrome_options()
                        
                        # Initialize WebDriver with error handling
                        try:
                            driver = webdriver.Chrome(options=chrome_options)
                            log_debug("ChromeDriver initialized successfully")
                        except WebDriverException as e:
                            st.error("Failed to initialize ChromeDriver. Please ensure Chrome and ChromeDriver are properly installed.")
                            log_debug(f"ChromeDriver initialization error: {str(e)}")
                            raise
                        
                        try:
                            log_debug(f"Navigating to URL: {url}")
                            driver.get(url)
                            
                            # Wait for page to load
                            time.sleep(3)  # Simple wait - consider using WebDriverWait in production
                            
                            log_debug("Taking screenshot...")
                            driver.save_screenshot(image_path)
                            log_debug(f"Screenshot saved to: {image_path}")
                            
                            # Display the image
                            st.image(image_path, caption="Captured Screenshot")
                        except Exception as e:
                            st.error(f"Error capturing webpage: {str(e)}")
                            log_debug(f"Webpage capture error: {str(e)}")
                            raise
                        finally:
                            driver.quit()
                            log_debug("ChromeDriver closed")
                    except Exception as e:
                        st.error(f"Failed to capture screenshot: {str(e)}")
                        if DEBUG:
                            st.exception(e)
    
    if st.button("Generate Test Cases from Image"):
        with st.spinner("Generating test cases from image..."):
            try:
                log_debug(f"Generating test cases for image: {image_path}")
                
                if image_path==None:
                    image_path="captured_screenshot.png"
                # Prepare the prompt

                try:
                    messages = [
                        ChatMessage(
                            role="system",
                            blocks=[TextBlock(text="You are an expert test engineer who excels in creating test cases for web products")]
                        ),
                        ChatMessage(
                            role="user",
                            blocks=[
                                ImageBlock(path=image_path),
                                TextBlock(text="Create test cases from the following screenshot."),
                            ],
                        )
                    ]
                except:
                    st.warning("Please Specify Image")
                    exit()
                
                log_debug(f"Image prompt messages: {messages}")
                
                # Get LLM response
                llm = get_llm()
                #resp = llm.chat(messages)
                gen = llm.stream_chat(messages=messages)
                st.markdown("### Generated Test Cases")
                # Stream the output
                output_container = st.empty()
                full_response = ""
                
            
                for response in gen:
                    full_response += response.delta
                    output_container.markdown(full_response)
                    time.sleep(0.01)
                
                log_debug("Received response from LLM")
                
                st.session_state['llm_response'] = full_response
                # # Display the response
                # st.markdown("### Generated Test Cases")
                # st.markdown(resp.message.blocks[0].text)
                
                # Clean up temporary files
                if option == "Capture from URL" and os.path.exists(image_path):
                    os.remove(image_path)
                    log_debug(f"Removed temporary file: {image_path}")
            except Exception as e:
                st.error(f"Error generating test cases: {str(e)}")
                log_debug(f"Test case generation error: {str(e)}")
                if DEBUG:
                    st.exception(e)

# Tab 3: Test Data Generation
with tab3:
    st.header("Generate Synthetic Test Data")
    
    test_cases = st.text_area(
        "Enter your test cases to generate matching test data:",
        height=200,
        help="Test cases to generate synthetic test data for"
    )
    
    if not test_cases:
        test_cases=st.session_state['llm_response']
        log_debug(f"Test cases generation error: {test_cases}")
        
    if st.button("Generate Test Data"):
        if not test_cases:
            st.warning("Please enter test cases first or generate test cases from previous tabs")
        else:
            with st.spinner("Generating test data..."):
                try:
                    # Prepare the prompt
                    messages = [
                        ChatMessage(
                            role="system",
                            blocks=[TextBlock(text="You are an expert test data generator for QA. Based on the following test cases, generate a table of synthetic test data suitable for automated testing. Include edge cases and realistic values.")]
                        ),
                        ChatMessage(
                            role="user",
                            blocks=[TextBlock(text=test_cases)]
                        ),
                    ]
                    
                    log_debug(f"Test data prompt: {messages}")
                    
                    # Get LLM response
                    llm = get_llm()
                    #resp = llm.chat(messages)
                    gen = llm.stream_chat(messages=messages)
                    st.markdown("### Generated Test Data")
                    # Stream the output
                    output_container = st.empty()
                    full_response = ""
                
                    for response in gen:
                        full_response += response.delta
                        output_container.markdown(full_response)
                        time.sleep(0.01)
                    st.session_state['llm_response'] = full_response
                    log_debug("Received response from LLM")
                except Exception as e:
                    st.error(f"Error generating test data: {str(e)}")
                    log_debug(f"Test data generation error: {str(e)}")
with tab4:
    st.header("Generate Test Cases from Code Changes")
    repoPath = st.text_input(
        "Enter repo path",
        placeholder="Repository path hosting code changes"
    )
    globals()["embedProvider"] = st.sidebar.radio(
        "Select Embedding Provider:",
        ("Ollama", "OpenAI")
    )
    if embedProvider=="Ollama":
        Settings.embed_model=get_ollamaEmbeddings()
    else:
        if not os.getenv("OPENAI_API_KEY") and globals()["OPENAI_API_KEY"]=="":
            globals()["OPENAI_API_KEY"]=st.sidebar.text_input(
                "OpenAI API Key",
                placeholder="Please Enter your OpenAI API key",key="openai_key")
        Settings.embed_model=get_openaiEmbeddings()
    
    Settings.llm=get_llm() 
    chkout_ver = st.selectbox(
        "Select checkout version history from HEAD",
        [i for i in range(0,5)],
        )
    output_container = st.empty()
    if chkout_ver:
        if not repoPath:
            st.warning("Please enter a repository path first")
            exit()
        changes=check_code_changes(repoPath,since_commit=f"HEAD~{chkout_ver}",agentic=False)
        output_container.markdown(f"***Changes found in: {[change for change in changes]}***")
    if st.button("Check Changes"):
        if not repoPath:
            st.warning("Please enter a repository path first")
            exit()
        changes=check_code_changes(repoPath,since_commit=f"HEAD~{chkout_ver}")
        output_container.code(changes)
        #output_container.markdown(changes)
    if st.button("Generate Test Cases from Changes"):
        if not repoPath:
            st.warning("Please enter a repository path first")
            exit()
        st.markdown("### Generated Test Cases")
        output_container = st.empty()
        index=_load_init_index(repo_path=repoPath,embProvider=embedProvider)
        changes=check_code_changes(repoPath,since_commit=f"HEAD~{chkout_ver}",agentic=False)
        # Stream the output
        full_response = ""
        log_debug(changes)
        for change in changes:
            output_container.markdown(full_response)
            full_response += "\n"+f"***Generating Test Cases for: {change}***\n"
            for response in generate_test_cases(index=index,changes=change,streaming=True):
                full_response+=response
                output_container.markdown(full_response)
                time.sleep(0.01)
            
with tab5:
    st.header("Generate Unit Tests from Code Changes")
    if option=="Ollama":
        Settings.embed_model=get_ollamaEmbeddings()
    else:
        Settings.embed_model=get_openaiEmbeddings()
    
    Settings.llm=get_llm() 
    if st.button("Generate Unit Tests"):
        output_container = st.empty()
        if not repoPath:
            st.warning("Please enter a repository path first")
            exit()
        changes=check_code_changes(repoPath,since_commit=f"HEAD~{chkout_ver}",agentic=False)
        full_response=''
        for change in changes:
            output_container.markdown(f"***Generating unit test for: {[change for change in changes]}***")
            response=generate_tests(repo_path=repoPath,file_path=change,streaming=False)
            full_response+=f"***Generating unit test for: {change}***\n{response.text}\n---\n"
            output_container.markdown(full_response)
            #response=generate_tests(repo_path=repoPath,file_path=change,streaming=True)
            # full_response = ""
            # for txt in response:
            #     full_response += txt.delta
            #     output_container.markdown(full_response)
            #     time.sleep(0.01) 

                    
    
st.sidebar.markdown("---")
st.sidebar.markdown("## About")
st.sidebar.markdown("This app demonstrates POC for QA Framework, Created by Adel Nazmy.")
if DEBUG:
    st.sidebar.divider()
    st.sidebar.warning("DEBUG SECTION",icon="🔥")
